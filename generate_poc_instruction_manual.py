"""
Generate POC Instruction Manual for HOS-13:
Step-by-step guide to run each POC (POC1 through POC6) independently.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = r"d:\Data_Delimited\Family_OS\jira\HOS13"


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "2C3E50")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True


def add_code_block(doc, code, language=""):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_step(doc, step_num, text):
    """Add a numbered step."""
    p = doc.add_paragraph()
    run = p.add_run(f"Step {step_num}: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    p.add_run(text)
    return p


def add_note(doc, text):
    """Add a note/warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("NOTE: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(196, 120, 0)
    p.add_run(text)
    return p


def generate_manual():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # === TITLE PAGE ===
    for _ in range(5):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Family OS")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("POC Instruction Manual")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph("")

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Step-by-Step Guide to Run Each POC Independently")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("JIRA: HOS-13 | February 17, 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run("POC Environment:\n").font.size = Pt(10)
    details.add_run("React Native 0.81.5 | Expo SDK 54.0.33 | TypeScript 5.9.2\n").font.size = Pt(10)
    details.add_run("Android (arm64-v8a) | Physical Device via USB\n").font.size = Pt(10)
    details.add_run("Node.js 22+ | npm 10+").font.size = Pt(10)

    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Prerequisites & Common Setup",
        "2. POC1-Calendar: Calendar Sync + UI",
        "3. POC2-PDFViewer: PDF Rendering",
        "4. POC3-CameraOCR: Camera + OCR",
        "5. POC4-Encryption: react-native-quick-crypto (BLOCKED)",
        "6. POC5-WebSocket: WebSocket + Zustand",
        "7. POC6-NobleCiphers: @noble/ciphers Encryption (VALIDATED)",
        "8. Troubleshooting",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==========================================================
    # SECTION 1: PREREQUISITES
    # ==========================================================
    doc.add_heading("1. Prerequisites & Common Setup", level=1)

    doc.add_heading("1.1 Required Software", level=2)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ["Software", "Version", "Purpose"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    prereqs = [
        ("Node.js", "22.x or later", "JavaScript runtime"),
        ("npm", "10.x or later (comes with Node.js)", "Package manager"),
        ("Android Studio", "Latest (with SDK 34+)", "Android SDK, emulator, build tools"),
        ("JDK", "17 (bundled with Android Studio)", "Java compiler for Android builds"),
        ("ADB", "Part of Android SDK platform-tools", "USB debugging, port forwarding"),
        ("Python", "3.8+ (for document generation only)", "Generates Word report documents"),
        ("Physical Android Device", "Android 7+ (API 24+), arm64-v8a", "Testing target (USB debugging enabled)"),
    ]
    for i, (sw, ver, purpose) in enumerate(prereqs):
        table.rows[i + 1].cells[0].text = sw
        table.rows[i + 1].cells[1].text = ver
        table.rows[i + 1].cells[2].text = purpose
    style_table(table)

    doc.add_paragraph("")

    doc.add_heading("1.2 Android Device Setup", level=2)

    add_step(doc, 1, "Enable Developer Options on your Android device:")
    add_code_block(doc, "Settings > About Phone > Tap 'Build Number' 7 times")

    add_step(doc, 2, "Enable USB Debugging:")
    add_code_block(doc, "Settings > Developer Options > Enable 'USB Debugging'")

    add_step(doc, 3, "Connect device via USB cable and verify ADB connection:")
    add_code_block(doc, "adb devices\n# Should show your device ID with 'device' status")

    add_step(doc, 4, "Set up port forwarding for Metro bundler:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081")

    doc.add_paragraph("")

    doc.add_heading("1.3 Environment Variables (Windows)", level=2)

    doc.add_paragraph("Ensure these environment variables are set:")
    add_code_block(doc,
        "ANDROID_HOME = C:\\Users\\<username>\\AppData\\Local\\Android\\Sdk\n"
        "JAVA_HOME = C:\\Program Files\\Android\\Android Studio\\jbr\n\n"
        "PATH should include:\n"
        "  %ANDROID_HOME%\\platform-tools\n"
        "  %ANDROID_HOME%\\tools"
    )

    doc.add_paragraph("")

    doc.add_heading("1.4 Common Build Commands", level=2)

    doc.add_paragraph("All POCs follow a similar workflow. The key commands are:")

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    headers = ["Command", "Description"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    cmds = [
        ("npm install", "Install all dependencies from package.json"),
        ("npx expo prebuild --clean", "Generate native android/ and ios/ directories from Expo config"),
        ("npx expo run:android", "Build and install APK on connected device (includes Metro start)"),
        ("npx expo start --dev-client", "Start Metro bundler only (if APK already installed)"),
        ("adb reverse tcp:8081 tcp:8081", "Forward Metro port to device over USB"),
        ("npx expo run:android --device", "Build targeting a specific connected device"),
        ("gradlew.bat app:installDebug -PreactNativeArchitectures=arm64-v8a", "Build for 64-bit ARM only (use inside android/ folder)"),
    ]
    for i, (cmd, desc) in enumerate(cmds):
        table.rows[i + 1].cells[0].text = cmd
        table.rows[i + 1].cells[1].text = desc
    style_table(table)

    add_note(doc, "POC2, POC3, POC4, and POC6 require Development Builds (not Expo Go) because they use native modules. POC1 and POC5 can also run via Development Builds for consistency.")

    doc.add_page_break()

    # ==========================================================
    # SECTION 2: POC1 - Calendar
    # ==========================================================
    doc.add_heading("2. POC1-Calendar: Calendar Sync + UI", level=1)

    doc.add_heading("2.1 Overview", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC1-Calendar/"),
        ("Purpose", "Validate expo-calendar for device calendar access and react-native-calendars + react-native-big-calendar for UI rendering"),
        ("Key Libraries", "expo-calendar v15.0.8, react-native-calendars v1.1314.0, react-native-big-calendar v4.19.0, react-native-paper v5.15.0"),
        ("Expected Result", "3 tabs: Sync (read/write device calendar), Month View (colored dots), Week View (timeline with overlap)"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("2.2 Steps to Run", level=2)

    add_step(doc, 1, "Navigate to the POC1 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC1-Calendar")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")

    add_step(doc, 3, "Generate native project files:")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Connect your Android device via USB and verify:")
    add_code_block(doc, "adb devices")

    add_step(doc, 5, "Build and install on device:")
    add_code_block(doc, "npx expo run:android")

    add_step(doc, 6, "If the app is already installed and you just need Metro:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    doc.add_paragraph("")
    doc.add_heading("2.3 What to Test", level=2)

    tests = [
        "Sync Tab: Grant calendar permissions when prompted. Verify device calendars are listed. Create a test event and verify it appears in the device calendar app.",
        "Month View Tab: Verify the month calendar renders with colored dots per family member. Tap dates to see events for that day.",
        "Week View Tab: Verify the week/timeline view renders with overlapping event support. Swipe left/right to navigate between weeks.",
    ]
    for t in tests:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph("")
    doc.add_heading("2.4 Key Code Snippet: Calendar Permission & Event Read", level=2)

    add_code_block(doc,
        'import * as Calendar from "expo-calendar";\n\n'
        "// Request calendar permissions\n"
        "const { status } = await Calendar.requestCalendarPermissionsAsync();\n"
        'if (status !== "granted") {\n'
        '  console.log("Calendar permission denied");\n'
        "  return;\n"
        "}\n\n"
        "// List all device calendars\n"
        "const calendars = await Calendar.getCalendarsAsync(Calendar.EntityTypes.EVENT);\n"
        "console.log(`Found ${calendars.length} calendars`);\n\n"
        "// Read events from a date range\n"
        "const events = await Calendar.getEventsAsync(\n"
        "  [calendars[0].id],\n"
        "  startDate,\n"
        "  endDate\n"
        ");\n\n"
        "// Create a new event\n"
        "const eventId = await Calendar.createEventAsync(calendars[0].id, {\n"
        '  title: "Family Dinner",\n'
        "  startDate: new Date(2026, 1, 17, 18, 0),\n"
        "  endDate: new Date(2026, 1, 17, 19, 30),\n"
        '  timeZone: "Asia/Kolkata",\n'
        "});"
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 3: POC2 - PDF Viewer
    # ==========================================================
    doc.add_heading("3. POC2-PDFViewer: PDF Rendering", level=1)

    doc.add_heading("3.1 Overview", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC2-PDFViewer/"),
        ("Purpose", "Validate react-native-pdf for rendering PDF documents of various sizes in Document Vault"),
        ("Key Libraries", "react-native-pdf v7.0.3, react-native-blob-util v0.24.7, @config-plugins/react-native-pdf v12.0.0, @config-plugins/react-native-blob-util v12.0.0"),
        ("Expected Result", "PDF renders with pinch-to-zoom, multi-page scroll, load timing, and modal overlay preview"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("3.2 Steps to Run", level=2)

    add_step(doc, 1, "Navigate to the POC2 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC2-PDFViewer")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")

    add_step(doc, 3, "Generate native project files (required -- config plugins need prebuild):")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Build and install on device:")
    add_code_block(doc, "npx expo run:android")

    add_step(doc, 5, "If the app is already installed:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    doc.add_paragraph("")
    doc.add_heading("3.3 What to Test", level=2)

    tests = [
        "Tap 'Simple PDF (1 page)' -- verify it renders with load time logged.",
        "Tap 'W-9 Form (6 pages)' -- verify multi-page scrolling and pinch-to-zoom work.",
        "Tap 'Tax Instructions (100+ pages)' -- verify large PDF scrolls smoothly without crashes.",
        "Tap 'Document Vault Preview' -- verify modal overlay appears over the PDF.",
        "Check the log output for load timing measurements.",
    ]
    for t in tests:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph("")
    doc.add_heading("3.4 Key Code Snippet: PDF Rendering", level=2)

    add_code_block(doc,
        'import Pdf from "react-native-pdf";\n\n'
        "// Render a PDF from a URL\n"
        "<Pdf\n"
        '  source={{ uri: "https://example.com/document.pdf" }}\n'
        "  onLoadComplete={(numberOfPages, filePath) => {\n"
        "    console.log(`Loaded ${numberOfPages} pages`);\n"
        "  }}\n"
        "  onPageChanged={(page, numberOfPages) => {\n"
        "    console.log(`Page ${page} of ${numberOfPages}`);\n"
        "  }}\n"
        "  onError={(error) => {\n"
        '    console.log("PDF Error:", error);\n'
        "  }}\n"
        "  style={{ flex: 1 }}\n"
        "  enablePaging={false}\n"
        "  horizontal={false}\n"
        "/>"
    )

    add_note(doc, "react-native-pdf requires config plugins (@config-plugins/react-native-pdf and @config-plugins/react-native-blob-util) for Expo compatibility. These are listed in app.json plugins array and activated during prebuild.")

    doc.add_page_break()

    # ==========================================================
    # SECTION 4: POC3 - Camera + OCR
    # ==========================================================
    doc.add_heading("4. POC3-CameraOCR: Camera + OCR", level=1)

    doc.add_heading("4.1 Overview", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC3-CameraOCR/"),
        ("Purpose", "Validate expo-camera for photo capture and @react-native-ml-kit/text-recognition for on-device OCR text extraction"),
        ("Key Libraries", "@react-native-ml-kit/text-recognition v2.0.0, expo-camera v17.0.10, expo-image-picker v17.0.10, expo-media-library v18.2.1"),
        ("Expected Result", "Camera capture / gallery pick -> ML Kit OCR -> extracted text with block coordinates and timing"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("4.2 Steps to Run", level=2)

    add_step(doc, 1, "Navigate to the POC3 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC3-CameraOCR")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")

    add_step(doc, 3, "Generate native project files:")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Build and install on device:")
    add_code_block(doc, "npx expo run:android")

    add_step(doc, 5, "If the app is already installed:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    doc.add_paragraph("")
    doc.add_heading("4.3 What to Test", level=2)

    tests = [
        "Grant camera and photo permissions when prompted.",
        "Tap 'Take Photo' -- camera opens. Point at a document/receipt with text and capture.",
        "Tap 'Pick from Gallery' -- select an image with text from your gallery.",
        "After capture/pick, OCR runs automatically. Verify extracted text appears on the results screen.",
        "Check OCR timing (should be milliseconds), text block count, and character count in the log.",
        "Verify block-level detail shows coordinates (bounding box) for each text block.",
    ]
    for t in tests:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph("")
    doc.add_heading("4.4 Key Code Snippet: OCR Text Extraction", level=2)

    add_code_block(doc,
        'import TextRecognition from "@react-native-ml-kit/text-recognition";\n'
        'import { CameraView, useCameraPermissions } from "expo-camera";\n'
        'import * as ImagePicker from "expo-image-picker";\n\n'
        "// Capture photo with camera\n"
        "const photo = await cameraRef.current.takePictureAsync();\n\n"
        "// OR pick from gallery\n"
        "const result = await ImagePicker.launchImageLibraryAsync({\n"
        "  mediaTypes: ImagePicker.MediaTypeOptions.Images,\n"
        "  quality: 1,\n"
        "});\n\n"
        "// Run OCR on the image\n"
        "const startTime = Date.now();\n"
        "const ocrResult = await TextRecognition.recognize(imageUri);\n"
        "const elapsed = Date.now() - startTime;\n\n"
        "console.log(`OCR completed in ${elapsed}ms`);\n"
        "console.log(`Found ${ocrResult.blocks.length} text blocks`);\n"
        "console.log(`Full text: ${ocrResult.text}`);\n\n"
        "// Access block-level details\n"
        "ocrResult.blocks.forEach((block, i) => {\n"
        "  console.log(`Block ${i}: ${block.text}`);\n"
        "  console.log(`  Position: ${JSON.stringify(block.frame)}`);\n"
        "  console.log(`  Lines: ${block.lines.length}`);\n"
        "});"
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 5: POC4 - Encryption (BLOCKED)
    # ==========================================================
    doc.add_heading("5. POC4-Encryption: react-native-quick-crypto (BLOCKED)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("STATUS: BLOCKED -- Persistent Nitro Module PKCS1 initialization failure. See POC6 for the working encryption alternative.")
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_heading("5.1 Overview", level=2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC4-Encryption/"),
        ("Purpose", "Validate react-native-quick-crypto for AES-256-GCM encryption (BLOCKED -- tests cannot execute)"),
        ("Key Libraries", "react-native-quick-crypto v1.0.11, expo-secure-store v15.0.8, expo-build-properties v1.0.10"),
        ("Expected Result", "App crashes with 'TypeError: Cannot read property PKCS1 of undefined' on any crypto operation"),
        ("Recommendation", "Use POC6-NobleCiphers instead. react-native-quick-crypto's Nitro Module does not initialize correctly with Expo SDK 54 + React Native 0.81.5."),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("5.2 Steps to Run (for reference only -- tests will FAIL)", level=2)

    add_step(doc, 1, "Navigate to the POC4 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC4-Encryption")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")

    add_step(doc, 3, "Generate native project files:")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Build for arm64-v8a ONLY (to avoid CMake ninja loop on Windows):")
    add_code_block(doc,
        "cd android\n"
        "gradlew.bat app:installDebug -PreactNativeArchitectures=arm64-v8a -x lint -x test\n"
        "cd .."
    )

    add_note(doc, "Do NOT use 'npx expo run:android' directly on Windows -- it triggers a CMake/ninja infinite loop for armeabi-v7a. Always build with the arm64-v8a architecture flag.")

    add_step(doc, 5, "Start Metro bundler:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    add_step(doc, 6, "Open the app on your device. Any test button will trigger the PKCS1 error.")

    doc.add_paragraph("")
    doc.add_heading("5.3 Known Errors", level=2)

    doc.add_paragraph("Error 1: CMake/Ninja Build Loop (RESOLVED)", style='List Bullet')
    add_code_block(doc,
        "ninja: error: manifest 'build.ninja' still dirty after 100 tries\n\n"
        "FIX: Build for arm64-v8a only:\n"
        "gradlew.bat app:installDebug -PreactNativeArchitectures=arm64-v8a -x lint -x test"
    )

    doc.add_paragraph("Error 2: Nitro Module PKCS1 Failure (UNRESOLVED)", style='List Bullet')
    add_code_block(doc,
        "TypeError: Cannot read property 'PKCS1' of undefined\n\n"
        "This error persists despite all fix attempts:\n"
        "1. Added react-native-quick-crypto to app.json plugins\n"
        "2. Installed expo-build-properties v1.0.10\n"
        "3. Enabled Hermes JS engine\n"
        "4. Deleted android/.cxx, android/build, android/app/build\n"
        "5. Deleted entire android/ + npx expo prebuild --clean\n"
        "6. Multiple full native rebuilds\n"
        "7. Verified all dependencies present\n\n"
        "STATUS: UNRESOLVED. Use POC6 (@noble/ciphers) instead."
    )

    doc.add_paragraph("")
    doc.add_heading("5.4 Key Code Snippet (for reference -- does NOT work)", level=2)

    add_code_block(doc,
        'import QuickCrypto from "react-native-quick-crypto";\n'
        'import * as SecureStore from "expo-secure-store";\n\n'
        "const { Buffer } = QuickCrypto;\n\n"
        "// Generate random bytes\n"
        "const key = QuickCrypto.randomBytes(32);   // <-- CRASHES: PKCS1 undefined\n"
        "const iv = QuickCrypto.randomBytes(12);\n\n"
        "// AES-256-GCM encrypt\n"
        'const cipher = QuickCrypto.createCipheriv("aes-256-gcm", key, iv);\n'
        'let encrypted = cipher.update("Hello Family OS", "utf8", "hex");\n'
        'encrypted += cipher.final("hex");\n'
        "const authTag = cipher.getAuthTag();\n\n"
        "// Store key securely\n"
        'await SecureStore.setItemAsync("docVaultKey", key.toString("hex"));'
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 6: POC5 - WebSocket
    # ==========================================================
    doc.add_heading("6. POC5-WebSocket: WebSocket + Zustand", level=1)

    doc.add_heading("6.1 Overview", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC5-WebSocket/"),
        ("Purpose", "Validate React Native's built-in WebSocket API with Zustand v5 for real-time family sync"),
        ("Key Libraries", "Native WebSocket API (built-in), Zustand v5.0.11"),
        ("Expected Result", "Connect to echo servers, send/receive messages, JSON family updates, auto-reconnect, Zustand state management"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("6.2 Steps to Run", level=2)

    add_step(doc, 1, "Navigate to the POC5 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC5-WebSocket")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")

    add_step(doc, 3, "Generate native project files:")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Build and install on device:")
    add_code_block(doc, "npx expo run:android")

    add_step(doc, 5, "If the app is already installed:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    add_note(doc, "POC5 requires internet access on the device to connect to echo servers (wss://ws.postman-echo.com/raw or wss://echo.websocket.org). Ensure the device has WiFi or mobile data enabled.")

    doc.add_paragraph("")
    doc.add_heading("6.3 What to Test", level=2)

    tests = [
        "Tap 'Postman Echo' or 'WebSocket.org' to connect to an echo server. Status should change to 'Connected'.",
        "Type a message and tap 'Send' -- the echo server should return your message.",
        "Tap 'Send Family Update' -- sends a structured JSON message and receives it back.",
        "Disconnect WiFi briefly to test auto-reconnect. The app should reconnect automatically within 3 seconds.",
        "Check the message list shows sent (blue) and received (green) messages with timestamps.",
        "Verify reconnect count increments each time a reconnection occurs.",
    ]
    for t in tests:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph("")
    doc.add_heading("6.4 Key Code Snippet: WebSocket + Zustand Store", level=2)

    add_code_block(doc,
        'import { create } from "zustand";\n\n'
        "// Zustand store for WebSocket state\n"
        "const useWebSocketStore = create((set, get) => ({\n"
        '  status: "disconnected",\n'
        "  messages: [],\n"
        "  reconnectCount: 0,\n\n"
        "  connect: (serverUrl) => {\n"
        "    const ws = new WebSocket(serverUrl);\n\n"
        "    ws.onopen = () => {\n"
        '      set({ status: "connected" });\n'
        "    };\n\n"
        "    ws.onmessage = (event) => {\n"
        "      const message = {\n"
        "        id: Date.now().toString(),\n"
        "        text: event.data,\n"
        '        type: "received",\n'
        "        timestamp: new Date().toISOString(),\n"
        "      };\n"
        "      set((state) => ({ messages: [...state.messages, message] }));\n"
        "    };\n\n"
        "    ws.onclose = () => {\n"
        '      set({ status: "disconnected" });\n'
        "      // Auto-reconnect after 3 seconds\n"
        "      setTimeout(() => {\n"
        "        set((state) => ({ reconnectCount: state.reconnectCount + 1 }));\n"
        "        get().connect(serverUrl);\n"
        "      }, 3000);\n"
        "    };\n"
        "  },\n\n"
        "  sendMessage: (text) => {\n"
        "    wsRef.send(text);\n"
        "  },\n\n"
        "  sendFamilyUpdate: () => {\n"
        "    const update = JSON.stringify({\n"
        '      type: "family_update",\n'
        '      event: "task_completed",\n'
        '      member: "Parent",\n'
        '      task: "Pick up groceries",\n'
        "      timestamp: new Date().toISOString(),\n"
        "    });\n"
        "    wsRef.send(update);\n"
        "  },\n"
        "}));"
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 7: POC6 - NobleCiphers (VALIDATED)
    # ==========================================================
    doc.add_heading("7. POC6-NobleCiphers: @noble/ciphers Encryption (VALIDATED)", level=1)

    p = doc.add_paragraph()
    run = p.add_run("STATUS: ALL 5 TESTS PASSED -- Recommended encryption library for Family OS")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_heading("7.1 Overview", level=2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Field", "Details"),
        ("Folder", "POC6-NobleCiphers/"),
        ("Purpose", "Validate @noble/ciphers as working AES-256-GCM encryption library after POC4 was blocked. Executes the same 5 tests designed for POC4."),
        ("Key Libraries", "@noble/ciphers v1.3.0, expo-crypto v14.1.5, expo-secure-store v15.0.8"),
        ("Expected Result", "All 5 tests PASS: Random Bytes, AES-256-GCM Round-Trip, Wrong Key/Tamper Detection, Secure Store Integration, Performance Benchmark"),
        ("Special Requirement", "crypto-polyfill.ts MUST be imported before any @noble/ciphers code (Hermes engine lacks Web Crypto API)"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("7.2 Steps to Run", level=2)

    add_step(doc, 1, "Navigate to the POC6 directory:")
    add_code_block(doc, "cd d:\\Data_Delimited\\Family_OS\\jira\\HOS13\\POC6-NobleCiphers")

    add_step(doc, 2, "Install dependencies:")
    add_code_block(doc, "npm install")
    doc.add_paragraph("This installs @noble/ciphers, expo-crypto, and expo-secure-store among others.")

    add_step(doc, 3, "Generate native project files:")
    add_code_block(doc, "npx expo prebuild --clean")

    add_step(doc, 4, "Verify local.properties has the correct Android SDK path (if build fails):")
    add_code_block(doc,
        "# Check android/local.properties contains:\n"
        "sdk.dir=C\\:\\\\Users\\\\<username>\\\\AppData\\\\Local\\\\Android\\\\Sdk"
    )

    add_step(doc, 5, "Build and install on device:")
    add_code_block(doc, "npx expo run:android")

    add_step(doc, 6, "If the app is already installed, start Metro only:")
    add_code_block(doc, "adb reverse tcp:8081 tcp:8081\nnpx expo start --dev-client --port 8081")

    add_step(doc, 7, "Open the app on your device and tap 'Run All Tests'. All 5 tests should show green PASS checkmarks.")

    doc.add_paragraph("")
    doc.add_heading("7.3 What to Test", level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ["Test #", "Test Name", "What It Validates"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    test_details = [
        ("1", "Random Bytes Generation", "Generates 16, 32, 64-byte random values via expo-crypto polyfill. Verifies correct length and uniqueness. Validates CSPRNG works for key/nonce generation."),
        ("2", "AES-256-GCM Encrypt/Decrypt", "Encrypts test plaintext with 256-bit key + 12-byte nonce. Decrypts and verifies exact match. Core Document Vault encryption operation."),
        ("3", "Wrong Key / Tamper Detection", "Tests 3 failure scenarios: (a) wrong key, (b) tampered ciphertext, (c) wrong nonce. All must throw errors. Proves auth tag verification works (unlike Issue #798)."),
        ("4", "Secure Store Integration", "Full cycle: generate key -> store in expo-secure-store -> retrieve -> decrypt. Validates iOS Keychain / Android KeyStore integration."),
        ("5", "Performance Benchmark", "Measures encrypt+decrypt time for 100B, 1KB, 10KB, 100KB payloads. All should be sub-millisecond to low single-digit ms."),
    ]
    for i, (num, name, detail) in enumerate(test_details):
        table.rows[i + 1].cells[0].text = num
        table.rows[i + 1].cells[1].text = name
        table.rows[i + 1].cells[2].text = detail
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("7.4 Critical File: crypto-polyfill.ts", level=2)

    doc.add_paragraph(
        "This file is REQUIRED because React Native's Hermes JavaScript engine does not provide the "
        "Web Crypto API (crypto.getRandomValues) that @noble/ciphers needs. The polyfill uses expo-crypto "
        "(OS-level CSPRNG) and handles the 1024-byte-per-call limit by chunking."
    )

    add_code_block(doc,
        "// crypto-polyfill.ts\n"
        "// MUST be imported BEFORE any @noble/ciphers imports\n\n"
        'import * as ExpoCrypto from "expo-crypto";\n\n'
        'if (typeof globalThis.crypto === "undefined") {\n'
        "  (globalThis as any).crypto = {};\n"
        "}\n\n"
        'if (typeof globalThis.crypto.getRandomValues === "undefined") {\n'
        "  const MAX_CHUNK = 1024;\n"
        "  (globalThis.crypto as any).getRandomValues = <T extends ArrayBufferView>(\n"
        "    array: T\n"
        "  ): T => {\n"
        "    const target = new Uint8Array(\n"
        "      (array as unknown as Uint8Array).buffer,\n"
        "      (array as unknown as Uint8Array).byteOffset,\n"
        "      array.byteLength\n"
        "    );\n"
        "    let offset = 0;\n"
        "    while (offset < target.length) {\n"
        "      const chunkSize = Math.min(MAX_CHUNK, target.length - offset);\n"
        "      const bytes = ExpoCrypto.getRandomBytes(chunkSize);\n"
        "      target.set(\n"
        "        new Uint8Array(bytes.buffer, bytes.byteOffset, bytes.byteLength),\n"
        "        offset\n"
        "      );\n"
        "      offset += chunkSize;\n"
        "    }\n"
        "    return array;\n"
        "  };\n"
        "}"
    )

    doc.add_paragraph("")
    doc.add_heading("7.5 Entry Point: index.ts", level=2)

    doc.add_paragraph("The polyfill MUST be imported first in index.ts, before App or any other module:")

    add_code_block(doc,
        "// index.ts\n"
        'import "./crypto-polyfill";   // <-- MUST be first import\n'
        'import { registerRootComponent } from "expo";\n'
        'import App from "./App";\n\n'
        "registerRootComponent(App);"
    )

    add_note(doc, "ES module imports are hoisted, so the polyfill MUST be in a separate file imported first. Placing polyfill code directly in index.ts before other imports will NOT work because ES import hoisting moves all imports to the top regardless of code order.")

    doc.add_paragraph("")
    doc.add_heading("7.6 Key Code Snippet: AES-256-GCM Encrypt/Decrypt", level=2)

    add_code_block(doc,
        'import { gcm } from "@noble/ciphers/aes";\n'
        'import { randomBytes } from "@noble/ciphers/webcrypto";\n'
        'import * as SecureStore from "expo-secure-store";\n\n'
        "// Generate a 256-bit key and 12-byte nonce\n"
        "const key = randomBytes(32);     // 32 bytes = 256 bits\n"
        "const nonce = randomBytes(12);   // 12 bytes = 96 bits (GCM standard)\n\n"
        "// Encrypt\n"
        'const plaintext = new TextEncoder().encode("Hello, Family OS!");\n'
        "const aes = gcm(key, nonce);\n"
        "const ciphertext = aes.encrypt(plaintext);\n\n"
        "// Decrypt\n"
        "const aes2 = gcm(key, nonce);\n"
        "const decrypted = aes2.decrypt(ciphertext);\n"
        "const text = new TextDecoder().decode(decrypted);\n"
        'console.log(text);  // "Hello, Family OS!"\n\n'
        "// Wrong key detection (throws error)\n"
        "try {\n"
        "  const wrongKey = randomBytes(32);\n"
        "  gcm(wrongKey, nonce).decrypt(ciphertext);\n"
        '  console.log("ERROR: Should have thrown!");\n'
        "} catch (e) {\n"
        '  console.log("Correctly detected wrong key!");\n'
        "}\n\n"
        "// Store key in secure storage\n"
        "const keyHex = Array.from(key)\n"
        '  .map((b) => b.toString(16).padStart(2, "0"))\n'
        '  .join("");\n'
        'await SecureStore.setItemAsync("docVaultKey", keyHex);'
    )

    doc.add_page_break()

    # ==========================================================
    # SECTION 8: TROUBLESHOOTING
    # ==========================================================
    doc.add_heading("8. Troubleshooting", level=1)

    doc.add_heading("8.1 Common Issues & Solutions", level=2)

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    headers = ["Issue", "Cause", "Solution"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    issues = [
        (
            "'Unable to load script' / Metro not connected",
            "Metro bundler not running or device can't reach it over USB",
            "Run: adb reverse tcp:8081 tcp:8081\nThen: npx expo start --dev-client --port 8081\nEnsure USB debugging is enabled."
        ),
        (
            "Port 8081 already in use",
            "Another Metro/Node process is using port 8081",
            "Windows: taskkill /F /IM node.exe\nLinux/Mac: kill $(lsof -t -i:8081)\nThen restart Metro."
        ),
        (
            "'SDK location not found' during build",
            "android/local.properties missing or incorrect",
            "Create/edit android/local.properties:\nsdk.dir=C\\:\\\\Users\\\\<user>\\\\AppData\\\\Local\\\\Android\\\\Sdk"
        ),
        (
            "CMake ninja infinite loop (POC4 only)",
            "react-native-quick-crypto build loop on Windows for armeabi-v7a",
            "Build for arm64-v8a only:\ncd android\ngradlew.bat app:installDebug -PreactNativeArchitectures=arm64-v8a -x lint -x test"
        ),
        (
            "'crypto.getRandomValues must be defined' (POC6)",
            "crypto-polyfill.ts not imported before @noble/ciphers",
            "Ensure index.ts has: import './crypto-polyfill' as the FIRST import line, before registerRootComponent."
        ),
        (
            "expo-crypto getRandomBytes range error (POC6)",
            "expo-crypto has a 1024-byte limit per call",
            "Use the chunking polyfill in crypto-polyfill.ts (already implemented). Do NOT call ExpoCrypto.getRandomBytes() with values > 1024."
        ),
        (
            "'Cannot find module' errors after npm install",
            "Stale cache or incomplete install",
            "Delete node_modules and reinstall:\nrm -rf node_modules\nnpm install"
        ),
        (
            "Build fails with 'Execution failed for task :app:...'",
            "Stale native build cache",
            "Clean rebuild:\nnpx expo prebuild --clean\nnpx expo run:android"
        ),
        (
            "App installed but shows white screen",
            "Metro bundler disconnected or JS error",
            "Check Metro terminal for red error text. Shake device to open dev menu and check for errors. Run adb logcat | grep ReactNative for native errors."
        ),
    ]
    for i, (issue, cause, solution) in enumerate(issues):
        table.rows[i + 1].cells[0].text = issue
        table.rows[i + 1].cells[1].text = cause
        table.rows[i + 1].cells[2].text = solution
    style_table(table)

    doc.add_paragraph("")
    doc.add_heading("8.2 Clean Rebuild Procedure", level=2)

    doc.add_paragraph("If a POC is not building or behaving correctly, perform a clean rebuild:")

    add_code_block(doc,
        "# From the POC directory (e.g., POC6-NobleCiphers/)\n\n"
        "# 1. Delete node_modules and reinstall\n"
        "rm -rf node_modules\n"
        "npm install\n\n"
        "# 2. Delete native directories and regenerate\n"
        "rm -rf android ios\n"
        "npx expo prebuild --clean\n\n"
        "# 3. Build fresh\n"
        "npx expo run:android\n\n"
        "# OR for POC4 (Windows arm64 only):\n"
        "cd android\n"
        "gradlew.bat app:installDebug -PreactNativeArchitectures=arm64-v8a -x lint -x test\n"
        "cd ..\n"
        "adb reverse tcp:8081 tcp:8081\n"
        "npx expo start --dev-client --port 8081"
    )

    doc.add_paragraph("")
    doc.add_heading("8.3 POC Quick Reference", level=2)

    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    headers = ["POC", "Status", "Native Modules?", "Special Build Steps?", "Internet Required?"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    quick_ref = [
        ("POC1-Calendar", "GO", "Yes (expo-calendar)", "No", "No"),
        ("POC2-PDFViewer", "GO", "Yes (react-native-pdf)", "No", "Yes (loads PDF URLs)"),
        ("POC3-CameraOCR", "GO", "Yes (ML Kit, Camera)", "No", "No (on-device OCR)"),
        ("POC4-Encryption", "BLOCKED", "Yes (quick-crypto)", "Yes (arm64-v8a flag)", "No"),
        ("POC5-WebSocket", "GO", "No", "No", "Yes (echo servers)"),
        ("POC6-NobleCiphers", "GO", "No (pure JS crypto)", "No", "No"),
    ]
    for i, (poc, status, native, special, internet) in enumerate(quick_ref):
        table.rows[i + 1].cells[0].text = poc
        table.rows[i + 1].cells[1].text = status
        table.rows[i + 1].cells[2].text = native
        table.rows[i + 1].cells[3].text = special
        table.rows[i + 1].cells[4].text = internet
    style_table(table)

    doc.add_paragraph("")

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("--- End of POC Instruction Manual ---").font.color.rgb = RGBColor(150, 150, 150)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated on February 17, 2026 | JIRA: HOS-13")
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Family OS POC/Spike Validation | React Native 0.81.5 | Expo SDK 54")
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(9)

    filepath = os.path.join(OUTPUT_DIR, "Family_OS_POC_Instruction_Manual.docx")
    doc.save(filepath)
    print(f"Saved: {filepath}")
    return filepath


if __name__ == "__main__":
    print("Generating POC Instruction Manual...")
    f = generate_manual()
    print(f"\nDone! File created: {f}")
